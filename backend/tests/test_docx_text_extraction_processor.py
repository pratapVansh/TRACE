import uuid
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document as DocxDocument

from app.models.document import Document
from app.models.document_version import DocumentVersion
from app.models.ingestion_job import IngestionJob
from app.services.document_processing_exceptions import DocxTextExtractionError
from app.services.processing_status import ProcessingStage
from app.services.processors.base import ProcessingContext
from app.services.processors.docx_text_extraction import DocxTextExtractionProcessor


def _build_docx(text: str) -> bytes:
    from io import BytesIO

    document = DocxDocument()
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def docx_bytes() -> bytes:
    return _build_docx("Processor test content")


@pytest.fixture
def processing_context() -> ProcessingContext:
    document_id = uuid.uuid4()
    version_id = uuid.uuid4()
    document = Document(
        id=document_id,
        title="Manual",
        original_filename="manual.docx",
        doc_type="manual",
        status="processing",
        uploaded_by=None,
        extra_metadata={},
    )
    version = DocumentVersion(
        id=version_id,
        document_id=document_id,
        version_no=1,
        storage_uri="documents/test/v1/manual.docx",
        checksum_sha256="abc123",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_extension="docx",
        file_size_bytes=1024,
        is_latest=True,
    )
    job = IngestionJob(
        id=uuid.uuid4(),
        document_id=document_id,
        status="processing",
        stage=ProcessingStage.PROCESSING.value,
    )
    return ProcessingContext(document=document, version=version, job=job)


@pytest.fixture
def mock_storage(docx_bytes: bytes) -> MagicMock:
    storage = MagicMock()
    storage.read.return_value = docx_bytes
    return storage


@pytest.fixture
def mock_repository() -> AsyncMock:
    repository = AsyncMock()
    repository.upsert_extracted_text.return_value = MagicMock()
    repository.update_document_version.return_value = MagicMock()
    repository.update_ingestion_job.return_value = MagicMock()
    return repository


@pytest.fixture
def processor(mock_storage: MagicMock, mock_repository: AsyncMock) -> DocxTextExtractionProcessor:
    return DocxTextExtractionProcessor(
        storage=mock_storage,
        document_repository=mock_repository,
    )


@pytest.mark.asyncio
async def test_docx_processor_extracts_and_persists_text(
    processor: DocxTextExtractionProcessor,
    mock_storage: MagicMock,
    mock_repository: AsyncMock,
    processing_context: ProcessingContext,
) -> None:
    await processor.process(processing_context)

    mock_storage.read.assert_called_once_with(processing_context.version.storage_uri)
    mock_repository.update_ingestion_job.assert_awaited_with(
        processing_context.job.id,
        stage=ProcessingStage.TEXT_EXTRACTION.value,
    )
    saved_payload = mock_repository.upsert_extracted_text.await_args.kwargs
    assert "Processor test content" in saved_payload["full_text"]
    assert saved_payload["extraction_method"] == "python-docx"
    assert saved_payload["requires_ocr"] is False


@pytest.mark.asyncio
async def test_docx_processor_skips_non_docx_documents(
    processor: DocxTextExtractionProcessor,
    mock_storage: MagicMock,
    mock_repository: AsyncMock,
    processing_context: ProcessingContext,
) -> None:
    processing_context.version.file_extension = "pdf"

    await processor.process(processing_context)

    mock_storage.read.assert_not_called()
    mock_repository.upsert_extracted_text.assert_not_awaited()


@pytest.mark.asyncio
async def test_docx_processor_raises_when_storage_read_fails(
    processor: DocxTextExtractionProcessor,
    mock_storage: MagicMock,
    processing_context: ProcessingContext,
) -> None:
    from app.core.storage.exceptions import StorageError

    mock_storage.read.side_effect = StorageError("missing")

    with pytest.raises(DocxTextExtractionError, match="Failed to read stored DOCX"):
        await processor.process(processing_context)
