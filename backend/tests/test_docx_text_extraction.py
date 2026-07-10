from io import BytesIO

import pytest
from docx import Document as DocxDocument

from app.services.document_processing_exceptions import DocxTextExtractionError
from app.services.docx_text_extraction import extract_docx_text


def _build_docx() -> bytes:
    document = DocxDocument()
    document.add_heading("Safety Manual", level=1)
    document.add_paragraph("This procedure covers valve inspection.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tag"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "P-101"
    table.cell(1, 1).text = "120 psi"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_docx_text_reads_headings_paragraphs_and_tables() -> None:
    result = extract_docx_text(_build_docx())

    assert result.block_count >= 3
    assert "Safety Manual" in result.full_text
    assert "valve inspection" in result.full_text
    assert "P-101" in result.full_text
    assert any(block.block_type == "heading" for block in result.blocks)
    assert any(block.block_type == "paragraph" for block in result.blocks)
    assert any(block.block_type == "table" for block in result.blocks)


def test_extract_docx_text_preserves_document_order() -> None:
    result = extract_docx_text(_build_docx())

    block_types = [block.block_type for block in result.blocks]
    assert block_types.index("heading") < block_types.index("paragraph")
    assert block_types.index("paragraph") < block_types.index("table")


def test_extract_docx_text_rejects_empty_bytes() -> None:
    with pytest.raises(DocxTextExtractionError, match="empty"):
        extract_docx_text(b"")


def test_extract_docx_text_rejects_invalid_bytes() -> None:
    with pytest.raises(DocxTextExtractionError, match="Failed to open DOCX"):
        extract_docx_text(b"not-a-docx")
