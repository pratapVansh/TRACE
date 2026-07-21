from datetime import datetime
from pathlib import Path
from zipfile import BadZipFile, ZipFile

import docx
from docx.oxml.ns import qn

from app.core.config import settings
from app.core.logging import logger
from app.models.document import Document
from app.models.document_version import DocumentVersion
from app.processing.base import BaseProcessor


class DocxProcessor(BaseProcessor):
    name = "docx_processor"
    supported_extensions = frozenset({"docx"})
    supported_mime_types = frozenset({
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    })

    async def extract_text(
        self,
        document: Document,
        version: DocumentVersion,
    ) -> str:
        file_path = self._resolve_path(version)
        logger.info("Opening DOCX document_id=%s path=%s", document.id, file_path)

        try:
            doc = docx.Document(file_path)
        except Exception as exc:
            logger.error("Failed to open DOCX document_id=%s: %s", document.id, exc)
            return ""

        parts: list[str] = []

        self._extract_headers_footers(doc, parts, extract_header=True)
        self._extract_body(doc, parts)
        self._extract_headers_footers(doc, parts, extract_header=False)

        return "\n".join(parts).strip()

    async def extract_metadata(
        self,
        document: Document,
        version: DocumentVersion,
    ) -> dict:
        file_path = self._resolve_path(version)

        try:
            doc = docx.Document(file_path)
        except Exception as exc:
            logger.error("Failed to open DOCX for metadata document_id=%s: %s", document.id, exc)
            return self._fallback_metadata(document, version, file_path)

        cp = doc.core_properties
        page_break_count = self._count_page_breaks(doc)
        heading_count = self._count_headings(doc)
        table_count = len(doc.tables)
        paragraph_count = len(doc.paragraphs)
        image_count = self._count_images(doc)

        created = self._format_datetime(cp.created)
        modified = self._format_datetime(cp.modified)

        result = {
            "title": cp.title or document.title,
            "subject": cp.subject or "",
            "author": cp.author or "",
            "keywords": cp.keywords or "",
            "category": cp.category or "",
            "comments": cp.comments or "",
            "created": created,
            "modified": modified,
            "last_modified_by": cp.last_modified_by or "",
            "revision": cp.revision,
            "language": cp.language or "",
            "paragraph_count": paragraph_count,
            "heading_count": heading_count,
            "table_count": table_count,
            "image_count": image_count,
            "page_break_count": page_break_count,
            "file_size": self._get_file_size(file_path),
            "extension": "docx",
            "mime_type": version.mime_type,
        }

        if image_count > 0:
            result["requires_image_processing"] = True
            logger.info("Images detected document_id=%s count=%d", document.id, image_count)

        logger.info(
            "Metadata extracted document_id=%s paras=%d headings=%d tables=%d images=%d",
            document.id,
            paragraph_count,
            heading_count,
            table_count,
            image_count,
        )
        return result

    async def validate(
        self,
        document: Document,
        version: DocumentVersion,
    ) -> list[str]:
        warnings: list[str] = []
        file_path = self._resolve_path(version)

        if self._is_encrypted_zip(file_path):
            warnings.append("Password-protected DOCX file")
            return warnings

        try:
            doc = docx.Document(file_path)
        except BadZipFile:
            warnings.append("Corrupted or invalid DOCX file (not a valid ZIP archive)")
            return warnings
        except Exception as exc:
            err_str = str(exc).lower()
            if "password" in err_str or "encrypted" in err_str:
                warnings.append("Password-protected DOCX file")
            else:
                warnings.append(f"Cannot open DOCX: {exc}")
            return warnings

        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            warnings.append("DOCX document is empty (no paragraphs or tables)")
            return warnings

        empty_content = True
        for p in doc.paragraphs:
            if p.text.strip():
                empty_content = False
                break
        if empty_content:
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            empty_content = False
                            break
                    if not empty_content:
                        break
                if not empty_content:
                    break
        if empty_content:
            warnings.append("DOCX document is effectively empty (no text content)")

        return warnings

    async def process(
        self,
        document: Document,
        version: DocumentVersion,
    ) -> "ProcessingResult":
        from app.processing.models import ProcessingResult

        warnings = await self.validate(document, version)
        if warnings:
            metadata = await self.extract_metadata(document, version)
            return ProcessingResult(
                success=True,
                document_id=document.id,
                extracted_text="",
                metadata=metadata,
                warnings=warnings,
            )

        metadata = await self.extract_metadata(document, version)
        text = await self.extract_text(document, version)

        warnings_list: list[str] = []
        if metadata.get("requires_image_processing"):
            warnings_list.append("Embedded images detected. Image processing may be required.")

        return ProcessingResult(
            success=True,
            document_id=document.id,
            extracted_text=text,
            metadata=metadata,
            warnings=warnings_list,
        )

    def _extract_headers_footers(
        self,
        doc: docx.Document,
        parts: list[str],
        extract_header: bool = True,
    ) -> None:
        for i, section in enumerate(doc.sections):
            element = section.header if extract_header else section.footer
            if not element:
                continue
            linked_to_previous = getattr(element, 'is_linked_to_previous', False)
            has_paras = any(p.text.strip() for p in element.paragraphs)
            if not linked_to_previous and has_paras:
                parts.append(f"[{label} {i + 1}]")
                for p in element.paragraphs:
                    text = p.text.strip()
                    if text:
                        parts.append(text)
                parts.append(f"[/{label} {i + 1}]")

    def _extract_body(self, doc: docx.Document, parts: list[str]) -> None:
        logger.info("Extracting paragraphs document_id=%s", id(doc))

        body_elements = doc.element.body
        table_index = 0

        for child in body_elements:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                paragraph = self._find_paragraph_for_element(doc, child)
                if paragraph is not None:
                    self._extract_paragraph(paragraph, parts)
            elif tag == "tbl":
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    self._extract_table(table, table_index, parts)
                    table_index += 1

    def _find_paragraph_for_element(
        self,
        doc: docx.Document,
        element,
    ) -> docx.text.paragraph.Paragraph | None:
        for p in doc.paragraphs:
            if p._p is element:
                return p
        return None

    def _extract_paragraph(
        self,
        paragraph: docx.text.paragraph.Paragraph,
        parts: list[str],
    ) -> None:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""

        is_heading = style_name.lower().startswith("heading")
        is_list = self._is_list_paragraph(paragraph)

        if is_heading:
            level = self._heading_level(style_name)
            prefix = "#" * level + " " if level > 0 else ""
            if text:
                parts.append(f"{prefix}{text}")
                parts.append("")
            return

        if is_list:
            list_info = self._list_info(paragraph)
            parts.append(f"  {list_info}{text}")
            parts.append("")
            return

        if text:
            parts.append(text)
            parts.append("")

    def _is_list_paragraph(
        self,
        paragraph: docx.text.paragraph.Paragraph,
    ) -> bool:
        num_pr = paragraph._p.find(qn("w:numPr"))
        return num_pr is not None

    def _list_info(
        self,
        paragraph: docx.text.paragraph.Paragraph,
    ) -> str:
        num_pr = paragraph._p.find(qn("w:numPr"))
        if num_pr is None:
            return "- "

        ilvl = num_pr.find(qn("w:ilvl"))
        level = int(ilvl.get(qn("w:val"))) if ilvl is not None else 0
        indent = "  " * level

        num_id_elem = num_pr.find(qn("w:numId"))
        if num_id_elem is not None:
            num_id = int(num_id_elem.get(qn("w:val")))
        else:
            num_id = 0

        is_ordered = self._is_numbered_list(num_id)

        if is_ordered:
            return f"{indent}1. "
        else:
            return f"{indent}- "

    def _is_numbered_list(self, num_id: int) -> bool:
        return num_id % 2 == 1

    def _heading_level(self, style_name: str) -> int:
        try:
            parts = style_name.split()
            if len(parts) >= 2 and parts[1].isdigit():
                return int(parts[1])
        except (IndexError, ValueError):
            pass
        return 1

    def _extract_table(
        self,
        table: docx.table.Table,
        index: int,
        parts: list[str],
    ) -> None:
        logger.info("Extracting table %d document_id=%s", index + 1, id(table))

        parts.append("")
        parts.append("-" * 40)
        parts.append(f"Table {index + 1}")
        parts.append("")

        rows = list(table.rows)
        if not rows:
            parts.append("(empty table)")
            parts.append("")
            parts.append("-" * 40)
            parts.append("")
            return

        col_widths: list[int] = []
        for row in rows:
            cells = [cell.text.strip() for cell in row.cells]
            for col_idx, cell_text in enumerate(cells):
                if col_idx >= len(col_widths):
                    col_widths.append(0)
                col_widths[col_idx] = max(col_widths[col_idx], len(cell_text))

        col_widths = [max(w, 3) for w in col_widths]

        separator = " | ".join("-" * w for w in col_widths)
        header_separator = "-" * len(separator)

        for row_idx, row in enumerate(rows):
            cells = [cell.text.strip() for cell in row.cells]
            padded = []
            for col_idx, cell_text in enumerate(cells):
                width = col_widths[col_idx] if col_idx < len(col_widths) else 20
                padded.append(cell_text.ljust(width))
            parts.append(" | ".join(padded))

            if row_idx == 0:
                parts.append(separator)
                parts.append(header_separator)

        parts.append("")
        parts.append("-" * 40)
        parts.append("")

    def _count_page_breaks(self, doc: docx.Document) -> int:
        count = 0
        for p in doc.paragraphs:
            for run in p.runs:
                for br in run._r.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        count += 1

        for p in doc.paragraphs:
            last_child = p._p[-1] if len(p._p) > 0 else None
            if last_child is not None:
                tag = last_child.tag.split("}")[-1] if "}" in last_child.tag else last_child.tag
                if tag == "lastRenderedPageBreak":
                    count += 1

        return count

    def _count_headings(self, doc: docx.Document) -> int:
        count = 0
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name.lower().startswith("heading"):
                count += 1
        return count

    def _count_images(self, doc: docx.Document) -> int:
        count = 0
        for p in doc.paragraphs:
            for run in p.runs:
                drawings = run._r.findall(qn("w:drawing"))
                count += len(drawings)

        count += len(doc.inline_shapes)

        for section in doc.sections:
            for element in (section.header, section.footer):
                if element is None:
                    continue
                for p in element.paragraphs:
                    for run in p.runs:
                        drawings = run._r.findall(qn("w:drawing"))
                        count += len(drawings)

        return count

    @staticmethod
    def _is_encrypted_zip(file_path: str) -> bool:
        try:
            with ZipFile(file_path, "r") as zf:
                names = [n.lower() for n in zf.namelist()]
                for name in names:
                    if "encryptioninfo" in name:
                        return True
            return False
        except Exception:
            return False

    @staticmethod
    def _format_datetime(dt: datetime | None) -> str | None:
        if dt is None:
            return None
        try:
            return dt.isoformat()
        except (ValueError, AttributeError):
            return None

    @staticmethod
    def _get_file_size(file_path: str) -> int | None:
        try:
            return Path(file_path).stat().st_size
        except OSError:
            return None

    @staticmethod
    def _fallback_metadata(
        document: Document,
        version: DocumentVersion,
        file_path: str,
    ) -> dict:
        return {
            "title": document.title,
            "author": None,
            "paragraph_count": 0,
            "heading_count": 0,
            "table_count": 0,
            "image_count": 0,
            "page_break_count": 0,
            "file_size": DocxProcessor._get_file_size(file_path),
            "extension": "docx",
            "mime_type": version.mime_type,
        }
